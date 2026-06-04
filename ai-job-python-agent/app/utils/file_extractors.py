import tempfile
from pathlib import Path

from fastapi import HTTPException, UploadFile


async def extract_resume_file(file: UploadFile) -> str:
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="empty resume file")

    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()
    if filename.endswith(".pdf") or content_type == "application/pdf":
        text = _extract_pdf(data)
    elif filename.endswith(".docx") or "wordprocessingml.document" in content_type:
        text = _extract_docx(data)
    elif filename.endswith(".txt") or content_type.startswith("text/"):
        text = data.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail="only PDF, DOCX and TXT are supported")

    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise HTTPException(status_code=400, detail="no text extracted from resume")
    return normalized


def _extract_pdf(data: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise HTTPException(status_code=400, detail="PyMuPDF is required for PDF parsing") from exc

    with fitz.open(stream=data, filetype="pdf") as document:
        return "\n".join(page.get_text() for page in document)


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise HTTPException(status_code=400, detail="python-docx is required for DOCX parsing") from exc

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp:
        temp.write(data)
        temp_path = Path(temp.name)
    try:
        document = docx.Document(str(temp_path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    finally:
        temp_path.unlink(missing_ok=True)
